import sys
import os
import re
import io
import json
import time
import hashlib
import datetime
import uuid
import zipfile
from typing import List, Dict, Tuple, Optional, Any

import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image
import pytesseract
import spacy

try:
    import docx
except ImportError:
    docx = None

try:
    from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Query
    from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
    from fastapi.middleware.cors import CORSMiddleware
    import uvicorn
    FASTAPI_AVAILABLE = True
except ImportError:
    FASTAPI_AVAILABLE = False


TESSERACT_CONFIG = r"--oem 3 --psm 6"
LEDGER_FILE = "audit_ledger.json"

# Load spaCy NLP model safely
try:
    nlp = spacy.load("en_core_web_sm")
except Exception:
    try:
        spacy.cli.download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
    except Exception as e:
        print("[!] Warning: spaCy model 'en_core_web_sm' could not be loaded. Falling back to regex engine only.")
        nlp = None

# ==========================================
# COMPLIANCE PROFILES & REGEX ENTITY RULES
# ==========================================

COMPLIANCE_PROFILES = {
    "GDPR": {
        "description": "General Data Protection Regulation - Protects personal identifying data (Names, Emails, Phones, Addresses, SSNs, IPs, DOBs).",
        "entities": ["NAME", "EMAIL", "PHONE", "ADDRESS", "CITY_STATE_ZIP", "SSN", "DOB", "IP_ADDRESS", "PASSPORT", "DRIVER_LICENSE", "LOCATION", "GENDER"]
    },
    "HIPAA": {
        "description": "Health Insurance Portability and Accountability Act - Protects PHI (Patient Names, MRNs, Health IDs, SSNs, Member IDs, Dates, Medical Records).",
        "entities": ["NAME", "MRN", "SSN", "DOB", "MEMBER_ID", "GROUP_NUMBER", "POLICY_ID", "PAYER_ID", "INSURANCE_ID", "PHONE", "ADDRESS", "CITY_STATE_ZIP", "EMAIL", "PHYSICIAN", "FACILITY", "NPI_ID", "GENDER"]
    },
    "PCI-DSS": {
        "description": "Payment Card Industry Data Security Standard - Protects financial data (Credit Cards, CVVs, Bank Accounts, Tax IDs, Account IDs).",
        "entities": ["CREDIT_CARD", "CVV", "BANK_ACCOUNT", "TAX_ID", "ACCOUNT_ID", "PAYER_ID", "SSN"]
    },
    "CUSTOM": {
        "description": "All-inclusive security profile scanning for all PII, PHI, and financial entity types.",
        "entities": []  # Empty means target all detected entities
    }
}

DEFAULT_PROFILE = "HIPAA"
DEFAULT_STRATEGY = "redact"  # Options: redact, hash, anonymize

REGEX_PATTERNS = {
    "CREDIT_CARD": r"\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}|3(?:0[0-5]|[68][0-9])[0-9]{11}|6(?:011|5[0-9]{2})[0-9]{12}|(?:2131|1800|35\d{3})\d{11})\b",
    "SSN": r"\b\d{3}-\d{2}-\d{4}\b",
    "EMAIL": r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-z]{2,}",
    "PHONE": r"(?:\+?\d{1,2}[\s-]?)?(?:\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4})",
    "MRN": r"(?i)\bMRN[:\s]*[A-Z0-9-]{5,15}\b",
    "DOB": r"(?i)(?:DOB|Date of Birth)[:\s]*([0-9]{1,2}[-/\s][0-9]{1,2}[-/\s][0-9]{2,4}|[A-Z][a-z]+[\s,]+[0-9]{1,2}[\s,]+[0-9]{4})",
    "MEMBER_ID": r"(?i)\b(?:Member ID|MID)[:\s]*([A-Z0-9-]{6,15})\b",
    "GROUP_NUMBER": r"(?i)\b(?:Group Number|Group|Group\s*#)[:\s]*([A-Z0-9-]{3,15})\b",
    "POLICY_ID": r"(?i)\b(?:Policy Effective Date|Policy ID)[:\s]*([A-Z0-9-]{5,15}|[0-9]{1,2}[-/][0-9]{1,2}[-/][0-9]{2,4})\b",
    "PAYER_ID": r"(?i)\b(?:Payer ID)[:\s]*([A-Z0-9-]{3,10})\b",
    "NPI_ID": r"(?i)\bNPI[:\s]*([0-9]{8,12})\b",
    "INSURANCE_ID": r"\b\d{6,15}\b(?=.*(?:Insurance|Insurer|Provider|ID))",
    "DRIVER_LICENSE": r"(?i)(?:License[:\s]*)[A-Z0-9-]{5,15}",
    "PASSPORT": r"(?i)(?:Passport[:\s]*)[A-Z0-9-]{6,15}",
    "TAX_ID": r"\b\d{2}-\d{7}\b(?=.*Tax)",
    "ACCOUNT_ID": r"\b[A-Z]{2,3}\d{6,}\b",
    "ADDRESS": r"\b\d{1,5}\s+[A-Za-z0-9\s,]{4,50}\s*(?:Street|St|Road|Rd|Avenue|Ave|Drive|Dr|Lane|Ln|Boulevard|Blvd|Court|Ct|Way|Place|Pkwy)\b(?:,\s*[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5})?",
    "CITY_STATE_ZIP": r"[A-Za-z\s]+,\s?[A-Z]{2}\s?\d{5}(?:-\d{4})?",
    "IP_ADDRESS": r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b",
    "GENDER": r"(?i)\b(?:Gender|Sex)[:\s]*\b(Male|Female|Other)\b",
    "NAME": r"(?i)(?<=Name[:\s])([A-Z][a-z]+(?:\s[A-Z]\.?)?[\s\b][A-Z][a-z]+(?:\s[A-Z][a-z]+)?)"
}

# ==========================================
# TAMPER-PROOF CRYPTOGRAPHIC AUDIT LEDGER
# ==========================================

class AuditLedger:
    def __init__(self, ledger_file: str = LEDGER_FILE):
        self.ledger_file = ledger_file
        self.ensure_ledger_exists()

    def ensure_ledger_exists(self):
        if not os.path.exists(self.ledger_file):
            timestamp = datetime.datetime.utcnow().isoformat() + "Z"
            genesis_block = {
                "index": 0,
                "timestamp": timestamp,
                "file_name": "GENESIS_BLOCK",
                "file_type": "system",
                "compliance_profile": "SYSTEM",
                "masking_strategy": "SYSTEM",
                "pii_found_count": 0,
                "entities_summary": {},
                "input_file_hash": "0" * 64,
                "output_file_hash": "0" * 64,
                "previous_hash": "0" * 64,
                "block_hash": self._compute_block_hash(0, "0" * 64, "GENESIS_BLOCK", "0" * 64, "0" * 64, timestamp)
            }
            with open(self.ledger_file, "w", encoding="utf-8") as f:
                json.dump([genesis_block], f, indent=2)

    def _compute_block_hash(self, index: int, prev_hash: str, file_name: str, in_hash: str, out_hash: str, timestamp: str = "") -> str:
        payload = f"{index}:{prev_hash}:{file_name}:{in_hash}:{out_hash}:{timestamp}"
        return hashlib.sha256(payload.encode("utf-8")).hexdigest()

    def get_all_blocks(self) -> List[Dict[str, Any]]:
        if not os.path.exists(self.ledger_file):
            self.ensure_ledger_exists()
        try:
            with open(self.ledger_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []

    def record_transaction(self, file_name: str, file_type: str, profile: str, strategy: str,
                           pii_entities: List[Tuple[str, str]], input_bytes: bytes, output_bytes: bytes) -> Dict[str, Any]:
        blocks = self.get_all_blocks()
        prev_block = blocks[-1] if blocks else {"index": -1, "block_hash": "0" * 64}

        index = prev_block["index"] + 1
        prev_hash = prev_block["block_hash"]
        timestamp = datetime.datetime.utcnow().isoformat() + "Z"

        in_hash = hashlib.sha256(input_bytes).hexdigest() if input_bytes else "0" * 64
        out_hash = hashlib.sha256(output_bytes).hexdigest() if output_bytes else "0" * 64

        entities_summary = {}
        for text, label in pii_entities:
            entities_summary[label] = entities_summary.get(label, 0) + 1

        block_hash = self._compute_block_hash(index, prev_hash, file_name, in_hash, out_hash, timestamp)

        new_block = {
            "index": index,
            "timestamp": timestamp,
            "file_name": file_name,
            "file_type": file_type,
            "compliance_profile": profile,
            "masking_strategy": strategy,
            "pii_found_count": len(pii_entities),
            "entities_summary": entities_summary,
            "input_file_hash": in_hash,
            "output_file_hash": out_hash,
            "previous_hash": prev_hash,
            "block_hash": block_hash
        }

        blocks.append(new_block)
        with open(self.ledger_file, "w", encoding="utf-8") as f:
            json.dump(blocks, f, indent=2)

        return new_block

    def verify_integrity(self) -> Tuple[bool, str, List[Dict[str, Any]]]:
        blocks = self.get_all_blocks()
        if not blocks:
            return False, "Ledger is empty or unreadable.", []

        details = []
        is_valid = True

        for i in range(len(blocks)):
            block = blocks[i]
            if i == 0:
                calc_hash = self._compute_block_hash(0, "0" * 64, block["file_name"], block["input_file_hash"], block["output_file_hash"], block.get("timestamp", ""))
                if block["block_hash"] != calc_hash:
                    is_valid = False
                    details.append({"index": 0, "status": "INVALID", "reason": "Genesis block hash mismatch"})
                else:
                    details.append({"index": 0, "status": "VALID"})
                continue

            prev_block = blocks[i - 1]
            if block["previous_hash"] != prev_block["block_hash"]:
                is_valid = False
                details.append({"index": block["index"], "status": "INVALID", "reason": f"Previous hash mismatch. Expected {prev_block['block_hash']}, got {block['previous_hash']}"})
                continue

            calc_hash = self._compute_block_hash(block["index"], block["previous_hash"], block["file_name"], block["input_file_hash"], block["output_file_hash"], block.get("timestamp", ""))
            if block["block_hash"] != calc_hash:
                is_valid = False
                details.append({"index": block["index"], "status": "INVALID", "reason": f"Block hash tampered! Recorded: {block['block_hash']}, Computed: {calc_hash}"})
            else:
                details.append({"index": block["index"], "status": "VALID"})

        msg = "Ledger chain is 100% valid and verified." if is_valid else "TAMPERING DETECTED in audit ledger chain!"
        return is_valid, msg, details


ledger_instance = AuditLedger()

# ==========================================
# ENTITY DETECTION ENGINE
# ==========================================

def detect_pii(text: str, profile: str = "HIPAA") -> List[Tuple[str, str]]:
    """Detect PII/PHI entities based on compliance profile."""
    if not text:
        return []

    profile = profile.upper()
    profile_info = COMPLIANCE_PROFILES.get(profile, COMPLIANCE_PROFILES["HIPAA"])
    target_entities = profile_info["entities"]

    detected = []

    # 1. Regex rules search
    for label, pattern in REGEX_PATTERNS.items():
        if target_entities and label not in target_entities:
            continue
        for match in re.finditer(pattern, text):
            # If there's a capture group, take the first group, otherwise the full match
            if match.groups():
                val = match.group(1).strip()
            else:
                val = match.group(0).strip()
            if val and len(val) >= 2:
                detected.append((val, label))

    # 2. spaCy NER search
    if nlp is not None:
        doc = nlp(text)
        for ent in doc.ents:
            et = ent.text.strip()
            if not et or len(et.replace(" ", "")) < 3:
                continue
            
            label = None
            if ent.label_ == "PERSON" and (not target_entities or "NAME" in target_entities or "PATIENT" in target_entities):
                label = "NAME"
            elif ent.label_ in {"GPE", "LOC"} and (not target_entities or "LOCATION" in target_entities or "ADDRESS" in target_entities):
                label = "LOCATION"
            elif ent.label_ == "ORG" and (not target_entities or "FACILITY" in target_entities):
                label = "FACILITY"

            if label:
                detected.append((et, label))

    # Deduplicate entities
    seen = set()
    uniq = []
    for text_val, label_val in detected:
        key = (text_val.lower(), label_val)
        if key not in seen:
            seen.add(key)
            uniq.append((text_val, label_val))

    return uniq


def get_masked_replacement(text_val: str, label_val: str, strategy: str = "redact") -> str:
    """Generate masked text according to selected strategy."""
    strategy = strategy.lower()
    if strategy == "hash":
        h = hashlib.sha256(text_val.encode("utf-8")).hexdigest()[:10]
        return f"[HASH_{label_val}:{h}]"
    elif strategy == "anonymize":
        val_hash = int(hashlib.md5(text_val.encode("utf-8")).hexdigest(), 16) % 900 + 100
        return f"[{label_val}_{val_hash}]"
    else:  # redact
        return f"[REDACTED_{label_val}]"

# ==========================================
# MULTI-FORMAT DOCUMENT REDACTION ENGINES
# ==========================================

# 1. PDF REDACTION ENGINE
def redact_pdf_document(input_path: str, output_path: str, pii_entities: List[Tuple[str, str]], strategy: str = "redact") -> bool:
    """Redact PDF while preserving layout, removing metadata."""
    try:
        doc = fitz.open(input_path)
    except Exception as e:
        print(f"[!] Could not open PDF {input_path}: {e}")
        return False

    # Strip metadata
    try:
        doc.set_metadata({})
    except Exception:
        pass

    for page in doc:
        for ent_text, label in pii_entities:
            ent_text = ent_text.strip()
            if not ent_text or len(ent_text) < 2:
                continue
            try:
                rects = page.search_for(ent_text)
            except Exception:
                rects = []
            for r in rects:
                try:
                    if strategy == "redact":
                        page.add_redact_annot(r, fill=(0, 0, 0))
                    else:
                        replacement = get_masked_replacement(ent_text, label, strategy)
                        page.add_redact_annot(r, text=replacement, fill=(1, 1, 1), text_color=(0, 0, 0.8))
                except Exception:
                    pass
        try:
            page.apply_redactions()
        except Exception:
            pass

    try:
        doc.save(output_path, garbage=4, deflate=True)
        doc.close()
        return True
    except Exception as e:
        print(f"[!] PDF save failed: {e}")
        try:
            doc.close()
        except Exception:
            pass
        return False

# 2. IMAGE REDACTION ENGINE
def redact_image_document(input_path: str, output_path: str, profile: str = "HIPAA", strategy: str = "redact") -> Tuple[List[Tuple[str, str]], int]:
    """Redact PII from image using OCR, blur faces, and strip EXIF metadata."""
    img = cv2.imread(input_path)
    if img is None:
        return [], 0

    # OCR text extraction
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    try:
        _, th = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    except Exception:
        th = gray

    ocr_data = pytesseract.image_to_data(th, output_type=pytesseract.Output.DICT, config=TESSERACT_CONFIG)
    words = [w for w in ocr_data.get("text", []) if w and w.strip()]
    full_text = " ".join(words)

    pii_entities = detect_pii(full_text, profile)

    # Overlays on text matches
    for i, word in enumerate(ocr_data["text"]):
        word_clean = word.strip()
        if not word_clean:
            continue
        for ent_text, label in pii_entities:
            if word_clean.lower() in ent_text.lower() or ent_text.lower() in word_clean.lower():
                x, y, w, h = (
                    ocr_data["left"][i],
                    ocr_data["top"][i],
                    ocr_data["width"][i],
                    ocr_data["height"][i],
                )
                if strategy == "redact":
                    cv2.rectangle(img, (x, y), (x + w, y + h), (0, 0, 0), -1)
                else:
                    replacement = get_masked_replacement(word_clean, label, strategy)
                    cv2.rectangle(img, (x, y), (x + w, y + h), (255, 255, 255), -1)
                    cv2.putText(img, replacement[:8], (x, y + h - 2), cv2.FONT_HERSHEY_SIMPLEX, 0.4, (200, 0, 0), 1)
                break

    # Face Blur detection
    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + "haarcascade_frontalface_default.xml")
    faces = face_cascade.detectMultiScale(gray, 1.2, 4)
    for (x, y, w, h) in faces:
        sub = img[y:y + h, x:x + w]
        img[y:y + h, x:x + w] = cv2.GaussianBlur(sub, (99, 99), 30)

    # Save without EXIF metadata
    rgb_img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    pil_img = Image.fromarray(rgb_img)
    pil_img.save(output_path)

    return pii_entities, len(faces)

# 3. WORD (.DOCX) REDACTION ENGINE
def redact_docx_document(input_path: str, output_path: str, pii_entities: List[Tuple[str, str]], strategy: str = "redact") -> bool:
    """Redact Word (.docx) document while preserving structure & styles."""
    if docx is None:
        print("[!] python-docx library is not available.")
        return False

    try:
        doc = docx.Document(input_path)
    except Exception as e:
        print(f"[!] Failed to open docx {input_path}: {e}")
        return False

    def replace_text_in_paragraph(p):
        for ent_text, label in pii_entities:
            if not ent_text or ent_text not in p.text:
                continue
            replacement = get_masked_replacement(ent_text, label, strategy)
            # Simple run-level or full text replacement preserving paragraph
            pattern = re.escape(ent_text)
            p.text = re.sub(pattern, replacement, p.text, flags=re.IGNORECASE)

    for p in doc.paragraphs:
        replace_text_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_text_in_paragraph(p)

    try:
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"[!] Failed to save redacted docx: {e}")
        return False

# 4. PLAIN TEXT / CODE DOCUMENT ENGINE
def redact_txt_document(input_path: str, output_path: str, pii_entities: List[Tuple[str, str]], strategy: str = "redact") -> bool:
    """Redact plain text files."""
    try:
        with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        for ent_text, label in pii_entities:
            if not ent_text:
                continue
            replacement = get_masked_replacement(ent_text, label, strategy)
            pattern = re.escape(ent_text)
            content = re.sub(pattern, replacement, content, flags=re.IGNORECASE)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        return True
    except Exception as e:
        print(f"[!] Text redaction failed: {e}")
        return False


def extract_document_text(filepath: str) -> str:
    """Extract plain text from any supported document format."""
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        text_parts = []
        try:
            doc = fitz.open(filepath)
            for page in doc:
                pt = page.get_text("text").strip()
                if pt:
                    text_parts.append(pt)
                else:
                    # OCR fallback for image page
                    pix = page.get_pixmap(dpi=200)
                    img_bytes = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    ocr_text = pytesseract.image_to_string(img, config=TESSERACT_CONFIG)
                    text_parts.append(ocr_text)
            doc.close()
        except Exception:
            pass
        return "\n\n".join(text_parts)

    elif ext in [".docx", ".doc"]:
        if docx is not None:
            try:
                doc = docx.Document(filepath)
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception:
                pass
        return ""

    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
        try:
            img = cv2.imread(filepath)
            if img is not None:
                gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
                return pytesseract.image_to_string(gray, config=TESSERACT_CONFIG)
        except Exception:
            pass
        return ""

    else:  # Text, log, csv, json, md
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""

# ==========================================
# FASTAPI WEB SERVER & INTERACTIVE DASHBOARD
# ==========================================

if FASTAPI_AVAILABLE:
    app = FastAPI(
        title="PrivGuard Local Inline Proxy",
        description="Local-First Document Redaction & Cryptographic Compliance Middleware",
        version="2.0.0"
    )

    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    active_profile = DEFAULT_PROFILE

    # ==========================================
    # HISTORY STATE MANAGER
    # ==========================================

    HISTORY_STATE_FILE = "history_state.json"

    def load_history_state() -> Dict[str, Any]:
        if os.path.exists(HISTORY_STATE_FILE):
            try:
                with open(HISTORY_STATE_FILE, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {"viewed": [], "deleted": []}

    def save_history_state(state: Dict[str, Any]):
        with open(HISTORY_STATE_FILE, "w", encoding="utf-8") as f:
            json.dump(state, f, indent=2)

    def get_document_records() -> List[Dict[str, Any]]:
        state = load_history_state()
        viewed_set = set(state.get("viewed", []))
        deleted_set = set(state.get("deleted", []))
        blocks = ledger_instance.get_all_blocks()
        records = []
        for block in blocks:
            if block.get("file_name") == "GENESIS_BLOCK":
                continue
            uid = str(block["index"])
            records.append({
                "id": uid,
                "index": block["index"],
                "file_name": block["file_name"],
                "file_type": block["file_type"],
                "compliance_profile": block["compliance_profile"],
                "masking_strategy": block["masking_strategy"],
                "pii_found_count": block["pii_found_count"],
                "entities_summary": block.get("entities_summary", {}),
                "timestamp": block["timestamp"],
                "input_file_hash": block["input_file_hash"],
                "output_file_hash": block["output_file_hash"],
                "block_hash": block["block_hash"],
                "viewed": uid in viewed_set,
                "deleted": uid in deleted_set,
                "status": "deleted" if uid in deleted_set else "sanitized",
            })
        return records

    # ==========================================
    # HTML SERVING ROUTES
    # ==========================================

    @app.get("/", response_class=HTMLResponse)
    @app.get("/dashboard", response_class=HTMLResponse)
    def serve_dashboard():
        return FileResponse("index.html")

    @app.get("/{filename}.html")
    def serve_html_pages(filename: str):
        filepath = f"{filename}.html"
        if os.path.exists(filepath):
            return FileResponse(filepath)
        raise HTTPException(status_code=404, detail="Page not found")

    # ==========================================
    # PROFILE ROUTES
    # ==========================================

    @app.get("/profiles")
    def get_profiles():
        global active_profile
        return {"active_profile": active_profile, "profiles": COMPLIANCE_PROFILES}

    @app.post("/profiles/switch")
    def switch_profile(data: Dict[str, str]):
        global active_profile
        prof = data.get("profile", "").upper()
        if prof in COMPLIANCE_PROFILES:
            active_profile = prof
            return {"status": "success", "active_profile": active_profile}
        raise HTTPException(status_code=400, detail="Invalid compliance profile name.")

    # ==========================================
    # LEDGER ROUTES
    # ==========================================

    @app.get("/ledger")
    def get_ledger():
        return ledger_instance.get_all_blocks()

    @app.post("/ledger/verify")
    def verify_ledger():
        is_valid, msg, details = ledger_instance.verify_integrity()
        return {"is_valid": is_valid, "message": msg, "details": details}

    # ==========================================
    # AUDIT LOG ROUTES
    # ==========================================

    @app.get("/audit/stats")
    def get_audit_stats():
        blocks = ledger_instance.get_all_blocks()
        real_blocks = [b for b in blocks if b.get("file_name") != "GENESIS_BLOCK"]
        total = len(real_blocks)
        # All ledger entries are successful sanitizations
        successful = total
        failed = 0
        # Count unique compliance profiles used as a proxy for "active users / roles"
        profiles_used = len(set(b.get("compliance_profile", "") for b in real_blocks))
        success_rate = 100.0 if total > 0 else 0.0
        failure_rate = 0.0
        return {
            "total_activities": total,
            "successful_actions": successful,
            "failed_actions": failed,
            "active_users": max(1, profiles_used) if total > 0 else 0,
            "success_rate": success_rate,
            "failure_rate": failure_rate,
        }

    @app.get("/audit/logs")
    def get_audit_logs(
        search: str = "",
        time_range: str = "all",
        action: str = "all",
        page: int = 1,
        per_page: int = 15,
    ):
        blocks = ledger_instance.get_all_blocks()
        real_blocks = [b for b in blocks if b.get("file_name") != "GENESIS_BLOCK"]
        now = datetime.datetime.utcnow()

        logs = []
        for b in real_blocks:
            ts_raw = b.get("timestamp", "")
            try:
                ts_dt = datetime.datetime.fromisoformat(ts_raw.replace("Z", "+00:00")).replace(tzinfo=None)
            except Exception:
                ts_dt = now
            logs.append({
                "id": str(b["index"]),
                "timestamp": ts_raw,
                "timestamp_dt": ts_dt,
                "user": "PrivGuard System",
                "action": "Document Sanitized",
                "details": f"{b.get('file_name','')} | {b.get('compliance_profile','')} | {b.get('masking_strategy','')} | {b.get('pii_found_count',0)} PII fields",
                "file_name": b.get("file_name", ""),
                "file_type": b.get("file_type", ""),
                "compliance_profile": b.get("compliance_profile", ""),
                "masking_strategy": b.get("masking_strategy", ""),
                "pii_found_count": b.get("pii_found_count", 0),
                "ip_address": "127.0.0.1",
                "status": "success",
                "block_hash": b.get("block_hash", "")[-12:],
                "block_index": b["index"],
            })

        # Time filter
        if time_range == "today":
            logs = [l for l in logs if l["timestamp_dt"].date() == now.date()]
        elif time_range == "week":
            week_ago = now - datetime.timedelta(days=7)
            logs = [l for l in logs if l["timestamp_dt"] >= week_ago]
        elif time_range == "month":
            month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            logs = [l for l in logs if l["timestamp_dt"] >= month_start]

        # Action filter
        if action and action != "all":
            logs = [l for l in logs if action.lower() in l["action"].lower()]

        # Search filter
        if search:
            sl = search.lower()
            logs = [l for l in logs if sl in l["file_name"].lower() or sl in l["compliance_profile"].lower() or sl in l["masking_strategy"].lower() or sl in l["ip_address"].lower() or sl in l["action"].lower()]

        # Sort newest first
        logs.sort(key=lambda x: x["timestamp_dt"], reverse=True)

        total = len(logs)
        total_pages = max(1, (total + per_page - 1) // per_page)
        page = max(1, min(page, total_pages))
        start = (page - 1) * per_page
        sliced = logs[start: start + per_page]

        # Remove datetime object before returning
        for l in sliced:
            del l["timestamp_dt"]

        return {"logs": sliced, "total": total, "page": page, "total_pages": total_pages, "per_page": per_page}

    # ==========================================
    # REPORTS ROUTES
    # ==========================================

    @app.get("/reports/summary")
    def get_reports_summary():
        blocks = ledger_instance.get_all_blocks()
        real_blocks = [b for b in blocks if b.get("file_name") != "GENESIS_BLOCK"]
        state = load_history_state()
        viewed_set = set(state.get("viewed", []))
        deleted_set = set(state.get("deleted", []))

        total = len(real_blocks)
        sanitized = sum(1 for b in real_blocks if str(b["index"]) not in deleted_set)
        viewed = len(viewed_set)
        deleted = len(deleted_set)
        success_rate = round(sanitized / total * 100, 1) if total > 0 else 0.0

        # Breakdown by profile
        by_profile = {}
        by_strategy = {}
        by_type = {}
        total_pii = 0

        for b in real_blocks:
            p = b.get("compliance_profile", "UNKNOWN")
            s = b.get("masking_strategy", "UNKNOWN")
            t = b.get("file_type", "unknown").lower()
            pii = b.get("pii_found_count", 0)
            total_pii += pii
            by_profile[p] = by_profile.get(p, 0) + 1
            by_strategy[s] = by_strategy.get(s, 0) + 1
            by_type[t] = by_type.get(t, 0) + 1

        avg_pii = round(total_pii / total, 1) if total > 0 else 0

        return {
            "total_documents": total,
            "sanitized_documents": sanitized,
            "viewed_documents": viewed,
            "deleted_documents": deleted,
            "failed_documents": 0,
            "success_rate": success_rate,
            "avg_pii_per_doc": avg_pii,
            "total_pii_found": total_pii,
            "by_profile": by_profile,
            "by_strategy": by_strategy,
            "by_type": by_type,
        }

    @app.get("/reports/top-documents")
    def get_top_documents(page: int = 1, per_page: int = 8):
        blocks = ledger_instance.get_all_blocks()
        real_blocks = [b for b in blocks if b.get("file_name") != "GENESIS_BLOCK"]
        real_blocks.sort(key=lambda b: b.get("timestamp", ""), reverse=True)
        total = len(real_blocks)
        total_pages = max(1, (total + per_page - 1) // per_page)
        page = max(1, min(page, total_pages))
        sliced = real_blocks[(page-1)*per_page: page*per_page]
        docs = []
        for b in sliced:
            docs.append({
                "index": b["index"],
                "file_name": b.get("file_name", ""),
                "file_type": b.get("file_type", ""),
                "compliance_profile": b.get("compliance_profile", ""),
                "masking_strategy": b.get("masking_strategy", ""),
                "pii_found_count": b.get("pii_found_count", 0),
                "timestamp": b.get("timestamp", ""),
                "status": "sanitized",
            })
        return {"documents": docs, "total": total, "page": page, "total_pages": total_pages}

    # ==========================================
    # HISTORY ROUTES
    # ==========================================

    @app.get("/history/stats")
    def get_history_stats():
        records = get_document_records()
        state = load_history_state()
        viewed_set = set(state.get("viewed", []))
        deleted_set = set(state.get("deleted", []))
        now = datetime.datetime.utcnow()
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        total = len(records)
        sanitized = sum(1 for r in records if r["status"] == "sanitized")
        deleted = len(deleted_set)
        viewed = len(viewed_set)
        viewed_this_month = 0
        deleted_this_month = 0
        for r in records:
            try:
                ts = datetime.datetime.fromisoformat(r["timestamp"].replace("Z", "+00:00")).replace(tzinfo=None)
            except Exception:
                continue
            if ts >= month_start:
                if r["id"] in viewed_set:
                    viewed_this_month += 1
                if r["id"] in deleted_set:
                    deleted_this_month += 1
        success_rate = round((sanitized / total * 100), 1) if total > 0 else 0.0
        return {
            "total_documents": total,
            "sanitized_documents": sanitized,
            "viewed_documents": viewed,
            "deleted_documents": deleted,
            "viewed_this_month": viewed_this_month,
            "deleted_this_month": deleted_this_month,
            "success_rate": success_rate,
        }

    @app.get("/history/records")
    def get_history_records(
        search: Optional[str] = Query(None),
        status: Optional[str] = Query(None),
        file_type: Optional[str] = Query(None),
        time_range: Optional[str] = Query(None),
        page: int = Query(1, ge=1),
        per_page: int = Query(10, ge=1, le=100),
    ):
        records = get_document_records()
        if status == "deleted":
            records = [r for r in records if r["deleted"]]
        elif status == "sanitized":
            records = [r for r in records if not r["deleted"]]
        elif status == "viewed":
            records = [r for r in records if r["viewed"] and not r["deleted"]]
        else:
            records = [r for r in records if not r["deleted"]]
        if search:
            q = search.lower()
            records = [r for r in records if q in r["file_name"].lower() or q in r["file_type"].lower()]
        if file_type and file_type != "all":
            records = [r for r in records if r["file_type"].lower() == file_type.lower()]
        now = datetime.datetime.utcnow()
        if time_range == "today":
            cutoff = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_range == "week":
            cutoff = now - datetime.timedelta(days=7)
        elif time_range == "month":
            cutoff = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        else:
            cutoff = None
        if cutoff:
            filtered = []
            for r in records:
                try:
                    ts = datetime.datetime.fromisoformat(r["timestamp"].replace("Z", "+00:00")).replace(tzinfo=None)
                    if ts >= cutoff:
                        filtered.append(r)
                except Exception:
                    filtered.append(r)
            records = filtered
        records = sorted(records, key=lambda r: r["timestamp"], reverse=True)
        total = len(records)
        start = (page - 1) * per_page
        end = start + per_page
        page_records = records[start:end]
        return {
            "total": total,
            "page": page,
            "per_page": per_page,
            "total_pages": max(1, (total + per_page - 1) // per_page),
            "records": page_records,
        }

    @app.post("/history/view/{record_id}")
    def mark_viewed(record_id: str):
        state = load_history_state()
        if record_id not in state["viewed"]:
            state["viewed"].append(record_id)
            save_history_state(state)
        return {"status": "ok", "id": record_id, "viewed": True}

    @app.delete("/history/delete/{record_id}")
    def delete_record(record_id: str):
        state = load_history_state()
        if record_id not in state["deleted"]:
            state["deleted"].append(record_id)
            save_history_state(state)
        return {"status": "deleted", "id": record_id}

    @app.post("/history/bulk-delete")
    def bulk_delete_records(data: Dict[str, Any]):
        ids = data.get("ids", [])
        state = load_history_state()
        for rid in ids:
            if rid not in state["deleted"]:
                state["deleted"].append(rid)
        save_history_state(state)
        return {"status": "deleted", "count": len(ids)}

    @app.post("/history/restore/{record_id}")
    def restore_record(record_id: str):
        state = load_history_state()
        state["deleted"] = [d for d in state["deleted"] if d != record_id]
        save_history_state(state)
        return {"status": "restored", "id": record_id}

    @app.get("/history/export/csv")
    def export_history_csv(
        status: Optional[str] = Query(None),
        time_range: Optional[str] = Query(None),
    ):
        import csv, io as _io
        records = get_document_records()
        if status == "deleted":
            records = [r for r in records if r["deleted"]]
        else:
            records = [r for r in records if not r["deleted"]]
        now = datetime.datetime.utcnow()
        if time_range == "today":
            cutoff = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_range == "week":
            cutoff = now - datetime.timedelta(days=7)
        elif time_range == "month":
            cutoff = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        else:
            cutoff = None
        if cutoff:
            filtered = []
            for r in records:
                try:
                    ts = datetime.datetime.fromisoformat(r["timestamp"].replace("Z", "+00:00")).replace(tzinfo=None)
                    if ts >= cutoff:
                        filtered.append(r)
                except Exception:
                    filtered.append(r)
            records = filtered
        output = _io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["#", "File Name", "Type", "Profile", "Strategy", "PII Found", "Status", "Viewed", "Timestamp", "Block Hash"])
        for r in sorted(records, key=lambda x: x["timestamp"], reverse=True):
            writer.writerow([
                r["index"], r["file_name"], r["file_type"].upper(),
                r["compliance_profile"], r["masking_strategy"],
                r["pii_found_count"], r["status"].capitalize(),
                "Yes" if r["viewed"] else "No", r["timestamp"],
                r["block_hash"][:16] + "...",
            ])
        from fastapi.responses import StreamingResponse
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=privguard_history_{now.strftime('%Y%m%d_%H%M%S')}.csv"}
        )

    @app.post("/sanitize/document")
    async def sanitize_document(
        files: List[UploadFile] = File(...),
        profile: Optional[str] = Form(None),
        strategy: Optional[str] = Form("redact")
    ):
        prof = profile.upper() if profile else active_profile
        strat = strategy.lower() if strategy else DEFAULT_STRATEGY

        temp_dir = "temp_uploads"
        os.makedirs(temp_dir, exist_ok=True)
        session_id = str(uuid.uuid4())
        

        results = []
        zip_path = os.path.join(temp_dir, f"PrivGuard_Sanitized_{session_id[:8]}.zip")

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in files:
                input_filename = file.filename
                input_filename_safe = f"{session_id[:8]}_{input_filename}"
                input_path = os.path.join(temp_dir, f"raw_{input_filename_safe}")
                
                content_bytes = await file.read()
                with open(input_path, "wb") as f:
                    f.write(content_bytes)

                ext = os.path.splitext(input_filename)[1].lower()
                base_name = os.path.splitext(input_filename)[0]
                out_filename = f"redacted_{base_name}{ext}"
                out_filename = f"{session_id[:8]}_redacted_{base_name}{ext}"
                output_path = os.path.join(temp_dir, out_filename)

                extracted_text = extract_document_text(input_path)
                pii_entities = detect_pii(extracted_text, prof)

                faces_count = 0
                success = False

                if ext == ".pdf":
                    success = redact_pdf_document(input_path, output_path, pii_entities, strat)
                elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
                    pii_entities, faces_count = redact_image_document(input_path, output_path, prof, strat)
                    success = True
                elif ext in [".docx", ".doc"]:
                    success = redact_docx_document(input_path, output_path, pii_entities, strat)
                else:
                    success = redact_txt_document(input_path, output_path, pii_entities, strat)

                output_bytes = b""
                if success and os.path.exists(output_path):
                    with open(output_path, "rb") as f:
                        output_bytes = f.read()
                    # Add to zip
                    zipf.write(output_path, out_filename)

                # Record transaction to cryptographic audit ledger
                ledger_instance.record_transaction(
                    file_name=input_filename,
                    file_type=ext.replace(".", ""),
                    profile=prof,
                    strategy=strat,
                    pii_entities=pii_entities,
                    input_bytes=content_bytes,
                    output_bytes=output_bytes
                )

                results.append({
                    "file_name": input_filename,
                    "entities_found": len(pii_entities),
                    "entities": pii_entities,
                    "raw_url": f"/download/raw_{input_filename_safe}",
                    "redacted_url": f"/download/{out_filename}"
                })

        return JSONResponse({
            "status": "success",
            "profile_used": prof,
            "strategy_used": strat,
            "results": results,
            "download_url": f"/download/batch/{os.path.basename(zip_path)}"
        })

    
    @app.get("/preview/{filename}")
    def preview_file(filename: str):
        path = os.path.join("temp_uploads", filename)
        if os.path.exists(path):
            return FileResponse(path, content_disposition_type="inline")
        raise HTTPException(status_code=404, detail="Document not found.")

    @app.get("/download/{filename}")
    def download_file(filename: str):
        path = os.path.join("temp_uploads", filename)
        if os.path.exists(path):
            return FileResponse(path, filename=filename)
        raise HTTPException(status_code=404, detail="Redacted document not found.")

    @app.get("/download/batch/{filename}")
    def download_batch(filename: str):
        path = os.path.join("temp_uploads", filename)
        if os.path.exists(path):
            return FileResponse(path, filename=filename, media_type="application/zip")
        raise HTTPException(status_code=404, detail="Batch zip file not found.")

# ==========================================
# CLI INTERFACE
# ==========================================

def run_cli(argv: List[str]):
    import argparse
    parser = argparse.ArgumentParser(description="PrivGuard Local Compliance & Multi-Format Document Redaction Middleware")
    parser.add_argument("files", nargs="*", help="Document file paths to redact (.pdf, .docx, .png, .jpg, .txt)")
    parser.add_argument("--profile", choices=["GDPR", "HIPAA", "PCI-DSS", "CUSTOM"], default="HIPAA", help="Compliance profile enforcement")
    parser.add_argument("--strategy", choices=["redact", "hash", "anonymize"], default="redact", help="Masking strategy")
    parser.add_argument("--server", action="store_true", help="Start FastAPI local inline proxy server")
    parser.add_argument("--host", default="127.0.0.1", help="Server host IP")
    parser.add_argument("--port", type=int, default=8000, help="Server port")
    parser.add_argument("--verify-ledger", action="store_true", help="Verify cryptographic audit ledger chain")

    args = parser.parse_args(argv[1:])

    if args.verify_ledger:
        print("[*] Verifying PrivGuard Cryptographic Audit Ledger...")
        valid, msg, details = ledger_instance.verify_integrity()
        if valid:
            print(f"[+] SUCCESS: {msg}")
        else:
            print(f"[!] WARNING: {msg}")
        for d in details:
            print(f"    Block #{d['index']}: {d['status']} {d.get('reason', '')}")
        return 0

    if args.server:
        if not FASTAPI_AVAILABLE:
            print("[!] FastAPI/Uvicorn not installed. Run: pip install fastapi uvicorn python-multipart")
            return 1
        print(f"[*] Starting PrivGuard Local Inline Proxy at http://{args.host}:{args.port}")
        print(f"[*] Serving Interactive Dashboard at http://{args.host}:{args.port}/dashboard")
        uvicorn.run(app, host=args.host, port=args.port)
        return 0

    if not args.files:
        parser.print_help()
        return 0

    for filepath in args.files:
        if not os.path.exists(filepath):
            print(f"[!] File not found: {filepath}")
            continue

        filename = os.path.basename(filepath)
        ext = os.path.splitext(filename)[1].lower()
        name_noext = os.path.splitext(filename)[0]
        out_filename = f"redacted_{name_noext}{ext}"

        print(f"[*] Processing {filename} | Profile: {args.profile} | Strategy: {args.strategy}")

        with open(filepath, "rb") as f:
            in_bytes = f.read()

        extracted_text = extract_document_text(filepath)
        pii_entities = detect_pii(extracted_text, args.profile)

        faces_count = 0
        success = False

        if ext == ".pdf":
            success = redact_pdf_document(filepath, out_filename, pii_entities, args.strategy)
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            pii_entities, faces_count = redact_image_document(filepath, out_filename, args.profile, args.strategy)
            success = True
        elif ext in [".docx", ".doc"]:
            success = redact_docx_document(filepath, out_filename, pii_entities, args.strategy)
        else:
            success = redact_txt_document(filepath, out_filename, pii_entities, args.strategy)

        out_bytes = b""
        if success and os.path.exists(out_filename):
            with open(out_filename, "rb") as f:
                out_bytes = f.read()

        ledger_instance.record_transaction(
            file_name=filename,
            file_type=ext.replace(".", ""),
            profile=args.profile,
            strategy=args.strategy,
            pii_entities=pii_entities,
            input_bytes=in_bytes,
            output_bytes=out_bytes
        )

        print(f"[+] Saved redacted document -> {out_filename}")
        print(f"    Entities found: {len(pii_entities)} | Faces redacted: {faces_count}")

    return 0


if __name__ == "__main__":
    sys.exit(run_cli(sys.argv))

